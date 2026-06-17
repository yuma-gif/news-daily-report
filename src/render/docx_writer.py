"""Word document rendering for structured news reports."""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt

from src.config import Settings
from src.llm.schemas import NewsItemOutput, NewsReportOutput
from src.render.image_handler import prepare_image

logger = logging.getLogger(__name__)

DOCX_IMAGE_WIDTH = Inches(6.0)
CHINESE_NUMERALS = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十", "十一", "十二"]


def write_docx_report(report: NewsReportOutput, settings: Settings) -> Path:
    output_dir = Path(settings.news_output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    timestamp = _timestamp_from_generated_at(report.generated_at)
    save_path = output_dir / f"News_Report_{timestamp}.docx"

    document = Document()
    _setup_document(document)
    _write_cover(document, report, settings.deepseek_model)

    for index, item in enumerate(_sorted_items(report.items), start=1):
        _write_news_item(document, index, item)

    document.save(save_path)
    return save_path


def _setup_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    style = document.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)


def _write_cover(document: Document, report: NewsReportOutput, model_name: str) -> None:
    document.add_heading("过去24小时全球重大新闻日报", level=1)

    subtitle_lines = [
        f"生成时间：{report.generated_at}",
        f"新闻时间范围：{report.time_range_start} 至 {report.time_range_end}",
        f"新闻数量：{report.total_items}",
        f"中国新闻数量：{report.china_items}",
        f"国际新闻数量：{report.international_items}",
        f"使用模型：{model_name}",
    ]
    for line in subtitle_lines:
        paragraph = document.add_paragraph(line)
        paragraph.paragraph_format.space_after = Pt(4)


def _write_news_item(document: Document, index: int, item: NewsItemOutput) -> None:
    numeral = CHINESE_NUMERALS[index - 1] if index <= len(CHINESE_NUMERALS) else str(index)
    document.add_heading(f"{numeral}、{item.title}", level=2)
    _add_chinese_title_translation(document, item)

    _add_section(document, "1. 新闻内容概要", item.summary)
    _add_section(document, "2. 详细新闻报道", item.detailed_report)
    _add_sources(document, item)
    _add_analysis(document, item)
    _add_images(document, item)


def _add_section(document: Document, title: str, body: str) -> None:
    _add_bold_paragraph(document, title)
    paragraph = document.add_paragraph(body)
    paragraph.paragraph_format.space_after = Pt(8)


def _add_chinese_title_translation(document: Document, item: NewsItemOutput) -> None:
    if not _looks_like_english_title(item.title):
        return

    translation = _first_sentence(item.summary)
    if not translation:
        return

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(f"中文译文：{translation}")
    run.italic = True


def _looks_like_english_title(title: str) -> bool:
    ascii_letters = sum(char.isalpha() and char.isascii() for char in title)
    chinese_chars = sum("\u4e00" <= char <= "\u9fff" for char in title)
    return ascii_letters >= 10 and ascii_letters > chinese_chars


def _first_sentence(text: str) -> str:
    stripped = text.strip()
    if not stripped:
        return ""
    for delimiter in ("。", "！", "？", ".", "!", "?"):
        if delimiter in stripped:
            return stripped.split(delimiter, 1)[0].strip() + delimiter
    return stripped[:120]


def _add_sources(document: Document, item: NewsItemOutput) -> None:
    _add_bold_paragraph(document, "来源链接：")
    for index, source in enumerate(item.sources, start=1):
        paragraph = document.add_paragraph(f"[{index}] {source.source_name} - {source.title} - {source.url}")
        paragraph.paragraph_format.space_after = Pt(3)


def _add_analysis(document: Document, item: NewsItemOutput) -> None:
    _add_bold_paragraph(document, "3. DeepSeek 分析")
    _add_bullet_group(document, "积极影响：", item.positive_impacts)
    _add_bullet_group(document, "消极影响：", item.negative_impacts)
    _add_bullet_group(document, "潜在风险：", item.risks)
    _add_bullet_group(document, "不确定性：", item.uncertainties)
    _add_bullet_group(document, "启发：", item.lessons)


def _add_images(document: Document, item: NewsItemOutput) -> None:
    _add_bold_paragraph(document, "4. 关键图片")
    images = item.images[:3]
    if not images:
        document.add_paragraph("本条新闻未找到可靠图片。")
        return

    inserted_count = 0
    for image in images:
        image_path = prepare_image(image)
        if image_path is None:
            continue
        try:
            document.add_picture(str(image_path), width=DOCX_IMAGE_WIDTH)
            inserted_count += 1
            caption = document.add_paragraph(
                f"图片说明：{image.caption}；来源：{image.source_name} - {image.source_url}"
            )
            caption.paragraph_format.space_after = Pt(8)
        except Exception as exc:
            logger.warning("Failed to insert image %s: %s", image.image_url, exc)

    if inserted_count == 0:
        document.add_paragraph("本条新闻未找到可靠图片。")


def _add_bullet_group(document: Document, title: str, values: list[str]) -> None:
    _add_bold_paragraph(document, title)
    if not values:
        document.add_paragraph("- 无")
        return
    for value in values:
        document.add_paragraph(f"- {value}")


def _add_bold_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.bold = True


def _sorted_items(items: list[NewsItemOutput]) -> list[NewsItemOutput]:
    return sorted(items, key=lambda item: (-item.importance_score, item.rank))


def _timestamp_from_generated_at(generated_at: str) -> str:
    digits = "".join(char for char in generated_at if char.isdigit())
    if len(digits) >= 12:
        return digits[:12]
    return digits.ljust(12, "0")
