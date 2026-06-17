import json

from docx import Document

from src.llm.schemas import NewsReportOutput
from src.render.docx_writer import write_docx_report


def test_docx_writer_generates_docx(tmp_path, make_settings):
    payload = json.loads(open("sample_report.json", encoding="utf-8").read())
    report = NewsReportOutput.model_validate(payload)
    settings = make_settings(tmp_path)

    output_path = write_docx_report(report, settings)

    assert output_path.exists()
    assert output_path.suffix == ".docx"
    document = Document(str(output_path))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "过去24小时全球重大新闻日报" in text
    assert "测试新闻标题" in text


def test_docx_writer_adds_chinese_translation_under_english_title(
    tmp_path,
    make_settings,
    make_report_payload,
):
    payload = make_report_payload()
    payload["items"][0]["title"] = "US and Iran exchange more strikes"
    payload["items"][0]["summary"] = "美伊连续第二天互相袭击，地区紧张局势继续升级。"
    report = NewsReportOutput.model_validate(payload)
    settings = make_settings(tmp_path)

    output_path = write_docx_report(report, settings)

    document = Document(str(output_path))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "中文译文：美伊连续第二天互相袭击，地区紧张局势继续升级。" in text
